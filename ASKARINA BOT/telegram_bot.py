# --- Import Library yang Diperlukan ---
from datetime import datetime
import os
import locale
import logging
from dotenv import load_dotenv
import pandas as pd
import tabulate
from openai import OpenAI
import google.generativeai as genai
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Application, CommandHandler, MessageHandler, filters,
    ContextTypes, ConversationHandler,
)
from docx import Document
from io import BytesIO
import gspread
from gspread.exceptions import WorksheetNotFound
from telegram.constants import ParseMode 
import telegram.error

# --- RAG Imports (CHROMA) ---
import chromadb
from chromadb.utils import embedding_functions 
import uuid 

# --- KONFIGURASI PENTING ---
# Ganti dengan nama model Telkom AI yang benar
TELKOM_MODEL_NAME = "telkom-ai-instruct" 
TELKOM_BASE_URL = "https://telkom-ai-dag-api.apilogy.id/Telkom-LLM/0.0.4" 

# --- Konfigurasi RAG Global ---
EMBEDDING_MODEL_NAME = 'all-MiniLM-L6-v2' 
RAG_COLLECTION_NAME = "dapros_telkom_ai"
RAG_CLIENT = None
DATABASE_DF = None # Variabel global untuk DataFrame

# --- Load .env ---
load_dotenv()

# --- Pengaturan Logging ---
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# --- Path Template dan Kredensial ---
SERVICE_ACCOUNT_FILE = "service_account.json"
SPREADSHEET_ID = os.getenv("SPREADSHEET_ID")
SHEET_NAME_TO_WRITE = "SPH Records" # Sheet untuk mencatat SPH
SHEET_ID_FOR_READ = os.getenv("SPREADSHEET_ID_FOR_READ", SPREADSHEET_ID)
SPREADSHEET_URL = f"https://docs.google.com/spreadsheets/d/{SHEET_ID_FOR_READ}/edit?usp=sharing"
SHEET_NAME_TO_READ = "Sheet1" # Sheet yang berisi data Internal (Dapros)

# Template paths berdasarkan pilihan
TEMPLATE_PATHS = {
    "HSI Only": "template/SPH [HSI ONLY].docx",
    "HSI + ANTARES": "template/SPH [HSI + ANTARES].docx",
    "HSI + NETMONK": "template/SPH [HSI + NETMONK].docx",
}

# --- API keys & clients ---
TELKOM_API_KEY = os.getenv("TELKOM_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")

TELKOM_CLIENT = None
if TELKOM_API_KEY and TELKOM_MODEL_NAME != "NAMA_MODEL_YANG_BENAR":
    try:
        TELKOM_CLIENT = OpenAI(
            api_key=TELKOM_API_KEY,
            base_url=TELKOM_BASE_URL + "/llm", 
            default_headers={"x-api-key": TELKOM_API_KEY},
        )
        logger.info(f"Telkom AI client diinisialisasi.")
    except Exception as e:
        logger.warning(f"Gagal inisialisasi Telkom AI client: {e}")

if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        GEMINI_CLIENT = genai.GenerativeModel('gemini-2.5-flash')
        logger.info("Gemini client diinisialisasi.")
    except Exception as e:
        logger.warning(f"Gagal inisialisasi Gemini client: {e}")
        GEMINI_CLIENT = None
else:
    GEMINI_CLIENT = None
    logger.warning("GEMINI_API_KEY tidak ditemukan.")

# --- States ---
MAIN_MENU, CHOOSE_MODE, HANDLE_QUERY, SPH_CHOOSE_TYPE, SPH_CUSTOMER = range(5)


# =========================================================================
# === HELPER FUNCTIONS (RAG, SPH, DB LOADING)
# =========================================================================

# --- Locale / tanggal ---
try:
    locale.setlocale(locale.LC_TIME, 'id_ID.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'id_ID')
    except:
        logger.warning("Locale id_ID tidak tersedia, menggunakan default sistem.")

def get_tanggal_sph():
    return datetime.now().strftime("%d %B %Y")

# --- Fungsi Helper SPH DOCX ---
def replace_placeholder(doc, data):
    def process_paragraph(paragraph, data):
        text = paragraph.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        if text != paragraph.text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(text)

    for p in doc.paragraphs:
        process_paragraph(p, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p, data)

def generate_sph_docx(data, template_path: str):
    try:
        if not os.path.isfile(template_path):
            logger.error(f"Template tidak ditemukan: {template_path}")
            return None
        document = Document(template_path)
        replace_placeholder(document, data)
        doc_io = BytesIO()
        document.save(doc_io)
        doc_io.seek(0)
        return doc_io
    except Exception as e:
        logger.exception(f"Error membuat SPH dari template: {e}")
        return None

# --- Fungsi Helper SPH Spreadsheet (BARU) ---
def record_sph_to_spreadsheet(data: dict) -> bool:
    """Mencatat data SPH yang dibuat ke Google Spreadsheet."""
    try:
        if not os.path.isfile(SERVICE_ACCOUNT_FILE) or not SPREADSHEET_ID:
            logger.error("Kredensial gspread tidak lengkap atau SPREADSHEET_ID tidak disetel.")
            return False
            
        gc = gspread.service_account(filename=SERVICE_ACCOUNT_FILE)
        sh = gc.open_by_key(SPREADSHEET_ID)
        
        try:
            ws = sh.worksheet(SHEET_NAME_TO_WRITE)
        except gspread.exceptions.WorksheetNotFound:
            ws = sh.add_worksheet(title=SHEET_NAME_TO_WRITE, rows="100", cols="10")
            headers = ["Tanggal SPH", "Nama Pelanggan", "Jenis SPH", "User Telegram ID", "Full Name User"]
            ws.append_row(headers)
            logger.info(f"Worksheet '{SHEET_NAME_TO_WRITE}' berhasil dibuat dengan header.")

        row_data = [
            data.get('Tanggal SPH', ''),
            data.get('Nama Pelanggan', ''),
            data.get('Jenis SPH', ''),
            data.get('User Telegram', ''),
            data.get('Full Name User', ''),
        ]
        
        ws.append_row(row_data)
        return True
        
    except Exception as e:
        logger.exception(f"Gagal mencatat SPH ke spreadsheet: {e}")
        return False

# --- RAG Indexing, Retrieval, dan Database Loader ---
def index_data_for_rag(df: pd.DataFrame):
    global RAG_CLIENT
    
    try:
        if df is None or df.empty:
            logger.warning("DataFrame kosong, RAG indexing dibatalkan.")
            return False

        if RAG_CLIENT is None:
            RAG_CLIENT = chromadb.PersistentClient(path="./rag_data_storage") 
            logger.info("RAG Client (ChromaDB) diinisialisasi dengan Persistence.")
            
        chunks = []
        for index, row in df.iterrows():
            # Pastikan kolom yang diambil ada di DataFrame
            nama_usaha = row.get('NAMA USAHA', 'N/A')
            ekosistem = row.get('EKOSISTEM', 'N/A')
            status_visit = row.get('STATUS VISIT', 'N/A')
            hasil_visit = row.get('HASIL VISIT', 'N/A')
            alamat = row.get('ALAMAT', 'N/A')
            nama_sales = row.get('NAMA SALES', 'N/A')
            keterangan = row.get('KET', 'N/A')
            no_sc_inet = row.get('NO SC/INET', 'N/A')
            
            customer_info = f"PELANGGAN: {nama_usaha}, EKOSISTEM: {ekosistem}, STATUS VISIT: {status_visit}, HASIL VISIT: {hasil_visit}, ALAMAT: {alamat}, SALES: {nama_sales}, KETERANGAN: {keterangan}, NO SC/INET: {no_sc_inet}"
            chunks.append(customer_info)

        sbert_ef = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name=EMBEDDING_MODEL_NAME
        )

        try:
            RAG_CLIENT.delete_collection(name=RAG_COLLECTION_NAME)
        except Exception:
            pass 

        collection = RAG_CLIENT.get_or_create_collection(
            name=RAG_COLLECTION_NAME,
            embedding_function=sbert_ef
        )

        ids = [str(uuid.uuid4()) for _ in chunks]
        collection.add(
            documents=chunks,
            ids=ids
        )
        
        logger.info(f"✅ RAG Indexing berhasil. {len(chunks)} dokumen terindeks.")
        return True
        
    except Exception as e:
        logger.exception(f"Gagal melakukan RAG Indexing: {e}")
        return False

def retrieve_relevant_data(query: str, n_results: int = 5) -> str:
    if RAG_CLIENT is None:
        return "ERROR: RAG system tidak terinisialisasi. Coba /reload_db."
        
    try:
        collection = RAG_CLIENT.get_collection(name=RAG_COLLECTION_NAME)
    except Exception:
        return "ERROR: RAG collection tidak ditemukan. Coba /reload_db."

    try:
        # Ambil minimal 15 hasil untuk konteks yang kaya
        results = collection.query(
            query_texts=[query],
            n_results=max(n_results, 15) 
        )
        
        context_docs = results.get('documents', [[]])[0]
        
        if not context_docs:
            return "Tidak ada data relevan ditemukan dalam database internal."

        return "\n---\n".join(context_docs)
    except Exception as e:
        logger.error(f"Error saat RAG retrieval: {e}")
        return "ERROR: Gagal saat melakukan pencarian di database internal."

def load_database_as_df(url=None, sheet_name=SHEET_NAME_TO_READ):
    if url is None:
        url = SPREADSHEET_URL
        
    logger.info("load_database_as_df: mulai memuat database")
    
    df = None
    # Prioritas 1: gspread (butuh service_account.json)
    if os.path.isfile(SERVICE_ACCOUNT_FILE) and SPREADSHEET_ID:
        try:
            gc = gspread.service_account(filename=SERVICE_ACCOUNT_FILE)
            sh = gc.open_by_key(SPREADSHEET_ID)
            ws = sh.worksheet(sheet_name)
            records = ws.get_all_records()
            df = pd.DataFrame(records)
            df = df.dropna(how='all')
            logger.info(f"Berhasil memuat database via gspread, baris: {len(df)}")
        except WorksheetNotFound:
             logger.error(f"Gagal memuat via gspread: Worksheet '{sheet_name}' tidak ditemukan.")
        except Exception as e:
             logger.exception(f"Gagal memuat via gspread: {e}")

    # Prioritas 2: Public URL (jika gspread gagal atau tidak dikonfigurasi)
    if df is None or df.empty:
        try:
            export_url = SPREADSHEET_URL
            if "pub?output=xlsx" not in SPREADSHEET_URL and "edit?usp=sharing" in SPREADSHEET_URL:
                 export_url = SPREADSHEET_URL.replace("/edit?usp=sharing", "/export?format=xlsx")
                 
            df = pd.read_excel(export_url, sheet_name=sheet_name, engine="openpyxl")
            df = df.dropna(how='all')
            logger.info(f"Berhasil memuat database dari publik URL, baris: {len(df)}")
        except Exception as e:
            logger.warning(f"Gagal memuat dari publik URL: {e}")

    if df is None or df.empty:
        logger.error("Semua metode gagal memuat database. Periksa konfigurasi.")
        return None
    
    # Panggil indexing di sini
    index_data_for_rag(df)
    
    return df

# Muat database saat startup
# NOTE: Dilakukan di fungsi main() agar bisa menggunakan job queue

# --- Helper untuk pesan panjang (SUDAH MEMILIKI FALLBACK PLAIN TEXT) ---
async def send_long_message(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str):
    MAX = 4000
    
    async def _send_part(message_text, part_num, total_parts):
        header = f"Bagian {part_num}/{total_parts}" if total_parts > 1 else ""
        text_to_send = f"**{header}**:\n{message_text}" if header else message_text
        
        try:
            # 1. Coba kirim dengan Markdown
            await update.message.reply_text(text_to_send, parse_mode=ParseMode.MARKDOWN)
        except telegram.error.BadRequest as e:
            # 2. Jika Markdown gagal, retry dengan Plain Text
            logger.warning(f"Markdown failed in send_long_message part {part_num}, trying plain text: {e}")
            await update.message.reply_text(text_to_send, parse_mode=None) # Plain Text
        except Exception as e:
            logger.error(f"Failed to send message part {part_num} in all formats: {e}")
            # Final fallback: send the raw text without any formatting attempt
            await update.message.reply_text(f"[ERROR: Gagal memproses format] {text_to_send}", parse_mode=None)

    if len(text) <= MAX:
        await _send_part(text, 1, 1)
        return
    
    # Split dan kirim
    parts = []
    cur = ""
    for line in text.splitlines(True):
        if len(cur) + len(line) > MAX:
            parts.append(cur)
            cur = line
        else:
            cur += line
    if cur:
        parts.append(cur)
        
    for i, p in enumerate(parts, 1):
        await _send_part(p, i, len(parts))

# =========================================================================
# === BOT HANDLERS (TELEGRAM INTERACTIONS)
# =========================================================================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Mengirim menu utama dan membersihkan data user."""
    context.user_data.clear()
    keyboard = [["Pilih Mode", "Buat SPH"]]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True)
    
    await update.message.reply_text(
        "Halo! Saya **ASKARINA**. Silakan pilih menu di bawah:", 
        reply_markup=reply_markup, 
        parse_mode=ParseMode.MARKDOWN
    )
    return MAIN_MENU

async def main_menu_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Menangani navigasi menu utama."""
    choice = update.message.text
    
    query_keyboard = [["Kembali ke Menu Utama"]]
    query_reply_markup = ReplyKeyboardMarkup(query_keyboard, resize_keyboard=True, one_time_keyboard=False)
    
    if choice == "Pilih Mode":
        keyboard = [["Data Internal", "Riset Prospek & Umum"], ["Kembali ke Menu Utama"]]
        await update.message.reply_text("Silakan pilih mode:", reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True))
        return CHOOSE_MODE
    
    if choice == "Data Internal":
        if TELKOM_CLIENT is None:
            await update.message.reply_text("⚠️ **Telkom AI Client tidak tersedia.** Periksa konfigurasi API Key dan `TELKOM_MODEL_NAME`.", parse_mode=ParseMode.MARKDOWN)
            return await start(update, context)
            
        context.user_data['mode'] = 'internal'
        await update.message.reply_text(
            "Mode diatur ke: **Data Internal**. Silakan ajukan pertanyaan.\n\n"
            "Anda dapat bertanya berkali-kali. Jika sudah selesai, tekan **Kembali ke Menu Utama**.",
            reply_markup=query_reply_markup,
            parse_mode=ParseMode.MARKDOWN
        )
        return HANDLE_QUERY
        
    if choice == "Riset Prospek & Umum":
        if GEMINI_CLIENT is None:
            await update.message.reply_text("⚠️ **Gemini Client tidak tersedia.** Periksa konfigurasi GEMINI_API_KEY.", parse_mode=ParseMode.MARKDOWN)
            return await start(update, context)
            
        context.user_data['mode'] = 'general'
        await update.message.reply_text(
            "Mode diatur ke: **Riset Prospek & Umum**. Silakan ajukan pertanyaan.\n\n"
            "Anda dapat bertanya berkali-kali. Jika sudah selesai, tekan **Kembali ke Menu Utama**.", 
            reply_markup=query_reply_markup,
            parse_mode=ParseMode.MARKDOWN
        )
        return HANDLE_QUERY
        
    if choice == "Kembali ke Menu Utama":
        return await start(update, context)
        
    if choice == "Buat SPH":
        context.user_data.setdefault('sph', {})
        keyboard = [["HSI Only"], ["HSI + ANTARES"], ["HSI + NETMONK"]]
        reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True)
        await update.message.reply_text("Pilih jenis SPH yang ingin dibuat:", reply_markup=reply_markup)
        return SPH_CHOOSE_TYPE
        
    return await start(update, context)

async def sph_choose_type(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Menangani pemilihan tipe SPH."""
    choice = update.message.text.strip()
    valid_types = TEMPLATE_PATHS.keys()
    
    if choice not in valid_types:
        await update.message.reply_text("Pilihan tidak valid. Silakan pilih dari tombol yang tersedia.")
        return SPH_CHOOSE_TYPE

    context.user_data['sph']['type'] = choice
    context.user_data['sph']['template_path'] = TEMPLATE_PATHS.get(choice)
    
    await update.message.reply_text(f"Anda memilih **{choice}**. Baik, siapa nama pelanggannya?", reply_markup=ReplyKeyboardRemove(), parse_mode=ParseMode.MARKDOWN)
    return SPH_CUSTOMER 

async def sph_get_customer(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Mendapatkan nama pelanggan SPH dan melanjutkan ke generate dokumen."""
    context.user_data.setdefault('sph', {})
    context.user_data['sph']['nama'] = update.message.text.strip()
 
    await update.message.reply_text("Data lengkap. Memproses dokumen SPH...")
    return await sph_generate_document(update, context)

async def sph_generate_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Membuat, mencatat, dan mengirim dokumen SPH."""
    context.user_data.setdefault('sph', {})
    sph_data = context.user_data['sph']
    
    sph_type = sph_data.get('type')
    template_path = sph_data.get('template_path')
    
    fixed_product = ""
    final_notes = "Penawaran ini tunduk pada syarat dan ketentuan yang berlaku. Penawaran berlaku 30 hari."

    if sph_type == "HSI Only":
        fixed_product = "High Speed Internet (HSI) dengan bandwidth sesuai permintaan."
    elif sph_type == "HSI + ANTARES":
        fixed_product = "High Speed Internet (HSI) dan solusi ANTARES (IoT Platform)."
    elif sph_type == "HSI + NETMONK":
        fixed_product = "High Speed Internet (HSI) dan solusi NETMONK (Network Monitoring)."
    else:
        await update.message.reply_text("⚠️ Jenis SPH tidak valid. Proses dibatalkan.")
        context.user_data.pop('sph', None)
        return await start(update, context)

    sph_data['tanggal'] = get_tanggal_sph()

    # Data untuk DOCX
    data_for_docx = {
        'nama': sph_data.get('nama', 'NAMA PELANGGAN'),
        'alamat': '...', 
        'produk': fixed_product, 
        'catatan_sph': final_notes, 
        'tanggal': sph_data.get('tanggal', ''),
    }

    # Generate Dokumen
    doc_io = generate_sph_docx(data_for_docx, template_path=template_path) 
    
    if not doc_io:
        await update.message.reply_text("Gagal membuat dokumen SPH. Periksa ketersediaan template file dan path.")
        context.user_data.pop('sph', None)
        return await start(update, context)

    # Record to Spreadsheet
    data_to_record = {
        'Tanggal SPH': sph_data.get('tanggal'),
        'Nama Pelanggan': sph_data.get('nama'),
        'Jenis SPH': sph_type,
        'User Telegram': update.effective_user.username or str(update.effective_user.id),
        'Full Name User': update.effective_user.full_name,
    }

    is_recorded = record_sph_to_spreadsheet(data_to_record)
    
    # Beri Feedback dan Kirim Dokumen
    if is_recorded:
        await update.message.reply_text("✅ Dokumen SPH berhasil dibuat dan **DATA TELAH DICATAT** ke spreadsheet.")
    else:
        await update.message.reply_text("✅ Dokumen SPH berhasil dibuat. ⚠️ **DATA GAGAL DICATAT** ke spreadsheet. Periksa log untuk detail.")

    filename = f"SPH_{sph_type.replace(' ', '_')}_{sph_data.get('nama','customer').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    await update.message.reply_document(document=doc_io, filename=filename)
    
    # Bersihkan dan Kembali
    context.user_data.pop('sph', None)
    return await start(update, context)


# --- FUNGSI HANDLE QUERY (MENGGUNAKAN RAG RETRIEVAL) ---
async def handle_query(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    global DATABASE_DF
    query = update.message.text.strip()
    mode = context.user_data.get('mode', 'general')
    answer = "Maaf, terjadi kesalahan saat memproses pertanyaan."
    
    # Handler untuk tombol "Kembali ke Menu Utama"
    if query == "Kembali ke Menu Utama":
        return await start(update, context)

    await update.message.reply_text(f"Mengolah pertanyaan Anda dalam mode: **{mode}**...", parse_mode=ParseMode.MARKDOWN)
    
    response_msg = await update.message.reply_text("Sedang memproses... ⏳")
    full_response = ""

    if mode == 'internal':
        if DATABASE_DF is None or DATABASE_DF.empty:
            full_response = "Database internal belum berhasil dimuat. Coba `/reload_db` dan pastikan **Sheet Name** dan **Spreadsheet ID** sudah benar."
        elif TELKOM_CLIENT is None:
            full_response = "Telkom AI client tidak tersedia. Pastikan `TELKOM_MODEL_NAME` dan API Key sudah benar."
        else:
            query_lower = query.lower()
            
            # --- QUERIES AGREGAT KHUSUS (Sales Names) ---
            if any(k in query_lower for k in ["siapa saja nama sales", "daftar sales", "semua sales", "sales mana"]):
                try:
                    sales_list = DATABASE_DF['NAMA SALES'].dropna().unique().tolist()
                    
                    if sales_list:
                        sales_text = "\n".join(f"- {name}" for name in sales_list)
                        answer = (
                            f"✅ Berdasarkan **seluruh data** yang tersedia ({len(DATABASE_DF)} baris), berikut adalah **Daftar Nama Sales** yang tercatat:\n"
                            f"\n{sales_text}\n"
                            f"\n*Total {len(sales_list)} nama sales ditemukan.*"
                        )
                    else:
                        answer = "Tidak ada nama sales yang tercatat dalam database internal."
                        
                    await context.bot.edit_message_text(
                        chat_id=update.effective_chat.id,
                        message_id=response_msg.message_id,
                        text=answer,
                        parse_mode=ParseMode.MARKDOWN
                    )
                    return HANDLE_QUERY
                    
                except KeyError:
                    full_response = "Kolom 'NAMA SALES' tidak ditemukan di database. Pastikan nama kolom sudah benar."
                except Exception as e:
                    full_response = f"Gagal mengambil daftar sales: `{e}`"
            # --- END QUERIES AGREGAT KHUSUS ---

            # --- RAG Logic (Jika bukan query khusus) ---
            if not full_response:
                try:
                    context_data = retrieve_relevant_data(query, n_results=20) 
                    
                    if context_data.startswith("ERROR"):
                        full_response = context_data
                    else:
                        full_prompt = (
                            f"Anda adalah ASKARINA, asisten data internal Telkom. Tugas utama Anda adalah **ekstraksi data yang akurat** dari 'DATA KONTEKS' yang diberikan.\n"
                            f"IKUTI ATURAN KETAT INI UNTUK MENJAMIN AKURASI:\n"
                            f"1. **ATURAN JAWABAN WAJIB AKURAT:** Anda **HARUS** menjawab *hanya* dan *eksklusif* berdasarkan informasi yang ada di 'DATA KONTEKS'. **JANGAN** pernah menambahkan informasi, opini, atau asumsi lain di luar konteks yang tersedia.\n"
                            f"2. **FORMAT:** Tampilkan hasil yang sudah difilter dalam format **daftar (list) berpoin** atau **paragraf deskriptif yang ringkas**. Jangan gunakan format tabel Markdown.\n"
                            f"3. **KEGAGALAN DATA:** Jika informasi di 'DATA KONTEKS' tidak cukup, nyatakan: 'Maaf, informasi yang relevan tidak ditemukan dalam database internal untuk pertanyaan ini.'\n\n"
                            f"DATA KONTEKS:\n{context_data}\n\nPertanyaan: {query}"
                        )

                        stream = TELKOM_CLIENT.chat.completions.create(
                            model=TELKOM_MODEL_NAME, 
                            messages=[{"role": "user", "content": full_prompt}],
                            stream=True
                        )
                        
                        for chunk in stream:
                            content = chunk.choices[0].delta.content
                            if content:
                                full_response += content
                                # Update pesan dengan streaming
                                if len(full_response) % 100 < len(content) or full_response.endswith((".", "\n", "?", "!")):
                                    if full_response.strip():
                                        try:
                                            await context.bot.edit_message_text(
                                                chat_id=update.effective_chat.id,
                                                message_id=response_msg.message_id,
                                                text=full_response + "▌",
                                                parse_mode=ParseMode.MARKDOWN
                                            )
                                        except telegram.error.BadRequest:
                                            try:
                                                await context.bot.edit_message_text(
                                                    chat_id=update.effective_chat.id,
                                                    message_id=response_msg.message_id,
                                                    text=full_response + "▌",
                                                    parse_mode=None
                                                )
                                            except Exception:
                                                pass
                                                
                        answer = full_response.strip()
                except Exception as e:
                    logger.exception(f"Error memproses query internal: {e}")
                    answer = f"Gagal memproses data internal dengan Telkom AI. Error: `{e}`"
    else:
        # general / riset (Gemini)
        if GEMINI_CLIENT:
            try:
                clean_query = query.encode('utf-8', 'ignore').decode('utf-8')
                
                stream = GEMINI_CLIENT.generate_content(
                    contents=[f"Anda adalah ASKARINA, asisten riset B2B. Jawab pertanyaan riset: {clean_query}"],
                    stream=True
                )
                
                for chunk in stream:
                    content = chunk.text
                    if content:
                        full_response += content
                        # Update pesan dengan streaming
                        if len(full_response) % 100 < len(content) or full_response.endswith((".", "\n", "?", "!")):
                            if full_response.strip():
                                try:
                                    await context.bot.edit_message_text(
                                        chat_id=update.effective_chat.id,
                                        message_id=response_msg.message_id,
                                        text=full_response + "▌",
                                        parse_mode=ParseMode.MARKDOWN
                                    )
                                except telegram.error.BadRequest:
                                    try:
                                        await context.bot.edit_message_text(
                                            chat_id=update.effective_chat.id,
                                            message_id=response_msg.message_id,
                                            text=full_response + "▌",
                                            parse_mode=None
                                        )
                                    except Exception:
                                        pass

                answer = full_response.strip()
            except Exception as e:
                logger.exception(f"Error memanggil Gemini: {e}")
                answer = f"Gagal melakukan riset dengan Gemini. Error: `{e}`"
        else:
            answer = "Gemini client tidak tersedia."

    # Final edit: ganti cursor dengan jawaban final
    if answer:
        try:
            # 1. Coba final edit dengan Markdown
            await context.bot.edit_message_text(
                chat_id=update.effective_chat.id,
                message_id=response_msg.message_id,
                text=answer,
                parse_mode=ParseMode.MARKDOWN
            )
        except telegram.error.BadRequest:
            # 2. Jika Markdown GAGAL, coba lagi di message yang sama dengan Plain Text
            try:
                await context.bot.edit_message_text(
                    chat_id=update.effective_chat.id,
                    message_id=response_msg.message_id,
                    text=answer,
                    parse_mode=None
                )
            except Exception as e:
                logger.error(f"Gagal final edit (Plain Text): {e}")
                # 3. Delete/send new message
                await context.bot.delete_message(
                    chat_id=update.effective_chat.id,
                    message_id=response_msg.message_id
                )
                await send_long_message(update, context, answer)
        except Exception:
            # Jika error lain, fallback ke send_long_message
            await send_long_message(update, context, answer)
    else:
        # Handle case where AI returns no response
        await context.bot.edit_message_text(
            chat_id=update.effective_chat.id,
            message_id=response_msg.message_id,
            text=answer or "Tidak ada respons dari AI.",
            parse_mode=ParseMode.MARKDOWN
        )

    return HANDLE_QUERY

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data.clear()
    await update.message.reply_text("Proses dibatalkan. Kembali ke menu utama.", reply_markup=ReplyKeyboardRemove())
    return await start(update, context)

async def reload_db(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text("Memuat ulang database dan mengindeks RAG... (cek log)")
    new_df = load_database_as_df()
    
    if new_df is None or new_df.empty:
        await update.message.reply_text(
            "Gagal memuat database. Periksa konfigurasi GSheet dan file `service_account.json`."
        )
    else:
        global DATABASE_DF
        DATABASE_DF = new_df
        await update.message.reply_text(f"Database berhasil dimuat ulang dan RAG Indexing selesai. Baris: {len(DATABASE_DF)}")

async def db_info(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if DATABASE_DF is None:
        await update.message.reply_text("Database belum ter-load atau kosong. Jalankan /reload_db.")
        return
    rows = len(DATABASE_DF)
    cols = DATABASE_DF.columns.tolist()
    
    rag_status = "TIDAK AKTIF"
    if RAG_CLIENT:
        try:
            collection = RAG_CLIENT.get_collection(name=RAG_COLLECTION_NAME)
            rag_status = f"AKTIF ({collection.count()} dokumen)"
        except Exception:
            rag_status = "AKTIF, tetapi koleksi tidak ditemukan."

    sample = tabulate.tabulate(DATABASE_DF.head(5), headers='keys', tablefmt='github', showindex=False)
    
    info_text = (
        f"**STATUS RAG:** {rag_status}\n"
        f"Baris Data: {rows}\n"
        f"Kolom: `{', '.join(cols)}`\n"
        f"\nContoh 5 baris:\n\n```\n{sample}\n```"
    )
    await update.message.reply_text(info_text, parse_mode=ParseMode.MARKDOWN)

# --- FUNGSI JOB QUEUE UNTUK REAL-TIME ---
def scheduled_db_reload(context: ContextTypes.DEFAULT_TYPE) -> None:
    """Memuat ulang database dan RAG index secara berkala di latar belakang."""
    
    logger.info("Jalankan scheduled_db_reload: Mulai memuat ulang database...")
    
    new_df = load_database_as_df() 
    
    if new_df is not None and not new_df.empty:
        global DATABASE_DF
        DATABASE_DF = new_df
        logger.info(f"✅ Scheduled reload berhasil. {len(DATABASE_DF)} baris terindeks ulang.")
    else:
        logger.error("⚠️ Scheduled reload gagal memuat data. DATABASE_DF tidak diperbarui.")
# ---------------------------------------------


def main():
    if not TELEGRAM_BOT_TOKEN:
        logger.error("TELEGRAM_BOT_TOKEN tidak ditemukan di .env")
        return
        
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    
    # Inisialisasi Database saat startup pertama kali
    global DATABASE_DF
    DATABASE_DF = load_database_as_df()

    # --- JobQueue untuk Real-Time Data (Reload setiap 5 menit) ---
    application.job_queue.run_repeating(
        scheduled_db_reload, 
        interval=300, 
        first=60, # Tunda 1 menit setelah startup agar tidak tabrakan dengan load pertama
        name='db_reloader'
    )
    logger.info("JobQueue untuk pemuatan data berkala telah diinisialisasi (5 menit).")
    # -------------------------------------

    # register command
    application.add_handler(CommandHandler("reload_db", reload_db))
    application.add_handler(CommandHandler("db_info", db_info))

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler(["start", "menu"], start)], 
        states={
            MAIN_MENU: [MessageHandler(filters.Regex(r'^(Pilih Mode|Buat SPH)$'), main_menu_handler)],
            CHOOSE_MODE: [MessageHandler(filters.Regex(r'^(Data Internal|Riset Prospek & Umum|Kembali ke Menu Utama)$'), main_menu_handler)],
            
            HANDLE_QUERY: [
                # Handler untuk tombol "Kembali ke Menu Utama"
                MessageHandler(filters.Regex(r'^Kembali ke Menu Utama$'), start), 
                CommandHandler(["start", "menu"], start),
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_query)
            ],
            
            SPH_CHOOSE_TYPE: [
                CommandHandler(["start", "menu"], start),
                MessageHandler(filters.TEXT & ~filters.COMMAND, sph_choose_type)
            ], 
            SPH_CUSTOMER: [
                CommandHandler(["start", "menu"], start),
                MessageHandler(filters.TEXT & ~filters.COMMAND, sph_get_customer)
            ], 
        },
        # Fallbacks
        fallbacks=[CommandHandler("cancel", cancel), CommandHandler("start", start), CommandHandler("menu", start)],
    )

    application.add_handler(conv_handler)
    logger.info("Bot sedang berjalan...")
    application.run_polling() 

if __name__ == "__main__":
    main()